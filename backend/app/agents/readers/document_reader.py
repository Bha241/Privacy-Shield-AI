from pathlib import Path
import io
import json
import logging
import os
from typing import Optional, List
import fitz  # PyMuPDF
from PIL import Image

logger = logging.getLogger(__name__)

# Keep PaddleX's automatic model downloads inside the project-owned writable
# cache instead of the machine-level home directory.
PADDLEX_CACHE_DIR = Path(__file__).resolve().parents[4] / ".local" / "paddlex"
PADDLEX_CACHE_DIR.mkdir(parents=True, exist_ok=True)
os.environ.setdefault("PADDLE_PDX_CACHE_HOME", str(PADDLEX_CACHE_DIR))
PADDLE_CACHE_DIR = PADDLEX_CACHE_DIR.parent / "paddle"
PADDLE_CACHE_DIR.mkdir(parents=True, exist_ok=True)
PADDLE_RUNTIME_HOME = PADDLEX_CACHE_DIR.parent / "paddle-home"
PADDLE_RUNTIME_HOME.mkdir(parents=True, exist_ok=True)
os.environ.setdefault("PADDLE_HOME", str(PADDLE_CACHE_DIR))
os.environ.setdefault("XDG_CACHE_HOME", str(PADDLEX_CACHE_DIR.parent))
os.environ.setdefault("PADDLE_EXTENSION_DIR", str(PADDLE_CACHE_DIR / "extensions"))

# PaddleOCR is the primary OCR engine. It is loaded lazily so the backend can
# still start when its model files are unavailable; EasyOCR remains the first
# runtime fallback in that case.
try:
    import importlib.util
    PADDLEOCR_AVAILABLE = importlib.util.find_spec("paddleocr") is not None
except Exception:
    PADDLEOCR_AVAILABLE = False

# Check OCR engines
try:
    import easyocr
    EASYOCR_AVAILABLE = True
except Exception:
    EASYOCR_AVAILABLE = False

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except Exception:
    PYTESSERACT_AVAILABLE = False


class DocumentReader:
    """Reads digital text from PDFs/documents and extracts text from images & scanned PDF pages via OCR."""

    def __init__(self, enable_ocr: bool = True, lang_list: List[str] = None, use_gpu: Optional[bool] = None):
        self.enable_ocr = enable_ocr
        self.lang_list = lang_list or ["en"]
        self._use_gpu = use_gpu
        self._paddle_reader = None
        self._easyocr_reader = None

    @property
    def paddle_reader(self):
        if self._paddle_reader is None and self.enable_ocr and PADDLEOCR_AVAILABLE:
            try:
                # PaddlePaddle 3.3 still resolves one legacy dataset cache
                # through USERPROFILE on Windows. Redirect it only while the
                # package initializes; the rest of the backend keeps its
                # normal process environment.
                previous_userprofile = os.environ.get("USERPROFILE")
                os.environ["USERPROFILE"] = str(PADDLE_RUNTIME_HOME)
                try:
                    from paddleocr import PaddleOCR
                    self._paddle_reader = PaddleOCR(
                        lang=self.lang_list[0],
                        use_doc_orientation_classify=False,
                        use_doc_unwarping=False,
                        use_textline_orientation=False,
                        enable_mkldnn=False,
                    )
                finally:
                    if previous_userprofile is None:
                        os.environ.pop("USERPROFILE", None)
                    else:
                        os.environ["USERPROFILE"] = previous_userprofile
                logger.info("Initialized PaddleOCR as the primary OCR engine")
            except Exception as e:
                logger.warning(f"PaddleOCR initialization warning; using EasyOCR fallback: {e}")
        return self._paddle_reader

    @property
    def easyocr_reader(self):
        if self._easyocr_reader is None and self.enable_ocr and EASYOCR_AVAILABLE:
            try:
                import easyocr
                import torch
                gpu_enabled = self._use_gpu if self._use_gpu is not None else torch.cuda.is_available()
                self._easyocr_reader = easyocr.Reader(self.lang_list, gpu=gpu_enabled)
                logger.info(f"Initialized EasyOCR (CUDA GPU Enabled: {gpu_enabled})")
            except Exception as e:
                logger.warning(f"EasyOCR initialization warning: {e}")
        return self._easyocr_reader

    def _ocr_image_bytes(self, image_bytes: bytes) -> str:
        """Run OCR on image bytes and return extracted text."""
        if not self.enable_ocr:
            return ""

        # 1. PaddleOCR (primary)
        if self.paddle_reader:
            try:
                import numpy as np
                image = np.asarray(Image.open(io.BytesIO(image_bytes)).convert("RGB"))
                results = self.paddle_reader.predict(
                    image,
                    use_doc_orientation_classify=False,
                    use_doc_unwarping=False,
                    use_textline_orientation=False,
                )
                texts = []
                for result in results:
                    payload = result
                    if hasattr(result, "json"):
                        payload = result.json
                        if callable(payload):
                            payload = payload()
                    if isinstance(payload, str):
                        payload = json.loads(payload)
                    if isinstance(payload, dict):
                        texts.extend(str(text).strip() for text in payload.get("rec_texts", []) if str(text).strip())
                if texts:
                    return "\n".join(texts)
            except Exception as e:
                logger.warning(f"PaddleOCR error; using EasyOCR fallback: {e}")

        # 2. EasyOCR fallback
        if self.easyocr_reader:
            try:
                results = self.easyocr_reader.readtext(image_bytes, detail=0)
                if results:
                    return "\n".join(results)
            except Exception as e:
                logger.warning(f"EasyOCR error: {e}")

        # 3. Optional Pytesseract fallback for installations without either
        # neural OCR engine or when both engines fail on a specific image.
        if PYTESSERACT_AVAILABLE:
            try:
                img = Image.open(io.BytesIO(image_bytes))
                return pytesseract.image_to_string(img).strip()
            except Exception as e:
                logger.warning(f"Pytesseract error: {e}")

        return ""

    def read_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file including digital text and image OCR.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text string from all pages
        """
        doc = fitz.open(file_path)
        full_text_blocks = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            digital_text = page.get_text("text").strip()

            ocr_text = ""
            images = page.get_images()

            # Run OCR if digital text is very short (<50 chars) or page contains embedded images
            if self.enable_ocr and (len(digital_text) < 50 or len(images) > 0):
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                ocr_text = self._ocr_image_bytes(img_bytes)

            combined_page_text = digital_text
            if ocr_text and ocr_text not in digital_text:
                combined_page_text += ("\n" if combined_page_text else "") + ocr_text

            if combined_page_text.strip():
                full_text_blocks.append(combined_page_text.strip())

        doc.close()
        return "\n\n".join(full_text_blocks)

    def read_image(self, file_path: str) -> str:
        """Extract text from standalone image file (PNG, JPG, JPEG, BMP)."""
        with open(file_path, "rb") as f:
            img_bytes = f.read()
        return self._ocr_image_bytes(img_bytes)

    def read_docx(self, file_path: str) -> str:
        """Extract text from Word .docx file."""
        # 1. Try python-docx if installed
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)
            if full_text:
                return "\n".join(full_text)
        except Exception as e:
            logger.warning(f"python-docx reader notice: {e}")

        # 2. Native ZipFile + XML parsing fallback for .docx
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            with zipfile.ZipFile(file_path) as z:
                xml_content = z.read("word/document.xml")
                root = ET.fromstring(xml_content)
                text_nodes = root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                text_list = [node.text for node in text_nodes if node.text]
                if text_list:
                    return "\n".join(text_list)
        except Exception as e:
            logger.warning(f"zipfile xml reader notice: {e}")

        # 3. Fallback text read
        try:
            return Path(file_path).read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""

    def read_doc(self, file_path: str) -> str:
        """Extract text from legacy Word .doc binary file."""
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            if text and text.strip():
                return text.strip()
        except Exception:
            pass

        try:
            with open(file_path, "rb") as f:
                content = f.read()
            import re
            printable = re.findall(rb'[\x20-\x7E\x09\x0A\x0D]{4,}', content)
            decoded = [p.decode('ascii', errors='ignore') for p in printable]
            clean_words = [w for w in decoded if not w.startswith(('WordDocument', 'Root Entry', 'CompObj', 'ObjectPool', 'SummaryInformation'))]
            return "\n".join(clean_words)
        except Exception:
            return ""

    def read_document(self, file_path: str) -> str:
        """
        Automatically detect file type and extract all text.

        Supports: .txt, .pdf, .docx, .doc, .png, .jpg, .jpeg, .bmp, .tiff
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")
        elif ext == ".pdf":
            return self.read_pdf(str(path))
        elif ext == ".docx":
            return self.read_docx(str(path))
        elif ext == ".doc":
            return self.read_doc(str(path))
        elif ext in {".png", ".jpg", ".jpeg", ".bmp", ".tiff"}:
            return self.read_image(str(path))
        else:
            try:
                return path.read_text(encoding="utf-8", errors="ignore")
            except Exception:
                return ""
